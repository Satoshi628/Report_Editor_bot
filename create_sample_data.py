"""サンプルの未完成週報(.docx)を作成するスクリプト。"""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK


def create_sample_draft():
    """指摘コメント付きのサンプル未完成週報を作成する。"""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)

    # ページ1: 最終原稿（指摘を反映済み）
    doc.add_paragraph("週報 2026年2月第1週（山田次郎）")
    doc.add_paragraph("")
    doc.add_paragraph("■ 今週の業務内容")
    doc.add_paragraph(
        "1. 顧客管理システムの改修\n"
        "   - 検索機能のレスポンス改善を実施し、平均応答時間を2.3秒から0.8秒に短縮\n"
        "   - SQLクエリのチューニングとインデックスの追加により達成\n"
        "   - 改善前後のベンチマーク結果を別紙にまとめた"
    )
    doc.add_paragraph(
        "2. 新規機能の要件定義\n"
        "   - ステークホルダー3名との要件ヒアリングを実施\n"
        "   - 要件定義書のドラフトを作成し、チームレビューを完了"
    )
    doc.add_paragraph("")
    doc.add_paragraph("■ 課題・問題点")
    doc.add_paragraph("- テスト環境の老朽化により、負荷テストの精度に懸念がある")
    doc.add_paragraph("")
    doc.add_paragraph("■ 来週の予定")
    doc.add_paragraph("- 顧客管理システム: 負荷テストの実施と結果分析")
    doc.add_paragraph("- 新規機能: 要件定義書の最終確定")

    # ページ区切り
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ページ2: 初版（指摘コメントの対象）
    doc.add_paragraph("週報 2026年2月第1週（山田次郎）")
    doc.add_paragraph("")
    doc.add_paragraph("■ 今週の業務")
    doc.add_paragraph(
        "1. 顧客管理システムの改修\n"
        "   - 検索を速くした\n"
        "   - SQLを直した"
    )
    doc.add_paragraph(
        "2. 新規機能の要件定義\n"
        "   - ヒアリングした\n"
        "   - 要件定義書を書いた"
    )
    doc.add_paragraph("")
    doc.add_paragraph("■ 問題")
    doc.add_paragraph("- テスト環境が古い")
    doc.add_paragraph("")
    doc.add_paragraph("■ 来週")
    doc.add_paragraph("- テストする")
    doc.add_paragraph("- 要件を決める")

    # 保存
    output_dir = Path(__file__).parent / "data" / "drafts"
    output_dir.mkdir(parents=True, exist_ok=True)
    filepath = output_dir / "draft_2026_02_w1.docx"
    doc.save(str(filepath))
    print(f"サンプル下書きを作成しました: {filepath}")


if __name__ == "__main__":
    create_sample_draft()
