"""週報データ管理サービス。

完成済み週報(.txt)と未完成週報(.docx)の読み込み・管理を行う。
"""

import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

import config


def get_completed_reports() -> list[dict]:
    """完成済み週報の一覧を取得する。

    Returns:
        list[dict]: 週報の一覧。各要素は id, filename, title を持つ。
    """
    reports = []
    completed_dir = config.COMPLETED_DIR
    if not completed_dir.exists():
        completed_dir.mkdir(parents=True, exist_ok=True)
        return reports

    for filepath in sorted(completed_dir.glob("*.txt")):
        first_line = ""
        with open(filepath, "r", encoding="utf-8") as f:
            first_line = f.readline().strip()

        reports.append({
            "id": filepath.stem,
            "filename": filepath.name,
            "title": first_line or filepath.stem,
        })

    return reports


def get_completed_report_content(report_id: str) -> Optional[dict]:
    """完成済み週報の内容を取得する。

    Args:
        report_id: 週報のID（ファイル名から拡張子を除いたもの）。

    Returns:
        dict: 週報の内容。見つからない場合は None。
    """
    filepath = config.COMPLETED_DIR / f"{report_id}.txt"
    if not filepath.exists():
        return None

    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.strip().split("\n")
    title = lines[0] if lines else report_id

    return {
        "id": report_id,
        "filename": filepath.name,
        "title": title,
        "content": content,
    }


def get_draft_reports() -> list[dict]:
    """未完成週報の一覧を取得する。

    Returns:
        list[dict]: 未完成週報の一覧。各要素は id, filename, page_count を持つ。
    """
    reports = []
    drafts_dir = config.DRAFTS_DIR
    if not drafts_dir.exists():
        drafts_dir.mkdir(parents=True, exist_ok=True)
        return reports

    for filepath in sorted(drafts_dir.glob("*.docx")):
        doc = Document(str(filepath))
        pages = _split_pages(doc)

        reports.append({
            "id": filepath.stem,
            "filename": filepath.name,
            "page_count": len(pages),
        })

    return reports


def get_draft_report_content(report_id: str) -> Optional[dict]:
    """未完成週報の内容を取得する。

    ページごとの本文と指摘コメントを構造化して返す。
    最後のページ = 初版、最初のページ = 最終原稿。

    Args:
        report_id: 週報のID。

    Returns:
        dict: 週報の構造化データ。見つからない場合は None。
    """
    filepath = config.DRAFTS_DIR / f"{report_id}.docx"
    if not filepath.exists():
        return None

    doc = Document(str(filepath))
    pages = _split_pages(doc)
    comments = _extract_comments(doc)

    page_data = []
    for i, page_text in enumerate(pages):
        page_comments = [c for c in comments if c.get("page_index") == i]
        page_data.append({
            "page_number": i + 1,
            "content": page_text,
            "comments": page_comments,
            "is_final": i == 0,
            "is_first_draft": i == len(pages) - 1,
        })

    return {
        "id": report_id,
        "filename": filepath.name,
        "page_count": len(pages),
        "pages": page_data,
    }


def get_all_completed_texts() -> list[str]:
    """全完成済み週報のテキストを取得する。

    Returns:
        list[str]: 各週報のテキスト内容のリスト。
    """
    texts = []
    completed_dir = config.COMPLETED_DIR
    if not completed_dir.exists():
        return texts

    for filepath in sorted(completed_dir.glob("*.txt")):
        with open(filepath, "r", encoding="utf-8") as f:
            texts.append(f.read())

    return texts


def _split_pages(doc: Document) -> list[str]:
    """Documentをページ区切りで分割する。

    Args:
        doc: python-docx Document オブジェクト。

    Returns:
        list[str]: ページごとのテキストリスト。
    """
    pages = []
    current_page_lines = []

    for paragraph in doc.paragraphs:
        # ページ区切りの検出
        has_page_break = False
        for run in paragraph.runs:
            if run._element.xml.find("w:br") != -1 and 'type="page"' in run._element.xml:
                has_page_break = True
                break

        if has_page_break and current_page_lines:
            pages.append("\n".join(current_page_lines))
            current_page_lines = []

        text = paragraph.text.strip()
        if text:
            current_page_lines.append(text)

    if current_page_lines:
        pages.append("\n".join(current_page_lines))

    return pages if pages else [""]


def _extract_comments(doc: Document) -> list[dict]:
    """Documentからコメントを抽出する。

    Args:
        doc: python-docx Document オブジェクト。

    Returns:
        list[dict]: コメントのリスト。各要素は author, text, page_index を持つ。
    """
    comments = []

    # コメントパーツの取得を試みる
    try:
        comment_part = None
        for rel in doc.part.rels.values():
            if "comments" in rel.reltype:
                comment_part = rel.target_part
                break

        if comment_part is None:
            return comments

        from lxml import etree

        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        root = etree.fromstring(comment_part.blob)

        for comment_elem in root.findall(".//w:comment", nsmap):
            author = comment_elem.get(f"{{{nsmap['w']}}}author", "不明")
            text_parts = []
            for p in comment_elem.findall(".//w:t", nsmap):
                if p.text:
                    text_parts.append(p.text)

            comments.append({
                "author": author,
                "text": "".join(text_parts),
                "page_index": 0,
            })

    except Exception:
        pass

    return comments
